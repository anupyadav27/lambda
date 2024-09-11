# question file diffrent 
import streamlit as st
from backend import process_text_and_user_input
import io
from docx import Document

# Path to the text file containing questions
QUESTIONS_FILE_PATH = 'questions.txt'

def load_questions():
    """Load questions from the text file."""
    try:
        with open(QUESTIONS_FILE_PATH, 'r') as file:
            questions = [line.strip() for line in file if line.strip()]
        return questions
    except FileNotFoundError:
        st.error("Questions file not found.")
        return []

def save_questions(questions):
    """Save questions to the text file."""
    with open(QUESTIONS_FILE_PATH, 'w') as file:
        for question in questions:
            file.write(f"{question}\n")

# Load questions from the file
questions = load_questions()

# Set up Streamlit interface
st.title("AWS Cloud Expert Analysis and Technical Documentation")

# User Input Form
st.subheader("Provide additional information:")

# Initialize session state
if 'current_question_index' not in st.session_state:
    st.session_state.current_question_index = 0
if 'user_responses' not in st.session_state:
    st.session_state.user_responses = {}
if 'final_document' not in st.session_state:
    st.session_state.final_document = None
if 'doc_buffer' not in st.session_state:
    st.session_state.doc_buffer = None

# Display answered questions
st.write("### Answered Questions:")
for q, a in st.session_state.user_responses.items():
    st.write(f"**{q}**")
    st.write(a)
    st.write("")

# Display the current question and handle form submission
if st.session_state.current_question_index < len(questions):
    current_question = questions[st.session_state.current_question_index]
    
    with st.form(key='user_input_form'):
        user_response = st.text_area(current_question, value=st.session_state.user_responses.get(current_question, ""))
        
        submit_button = st.form_submit_button("Next" if st.session_state.current_question_index < len(questions) - 1 else "Submit")
        
        # Add validation to ensure input is not empty
        if submit_button:
            if not user_response.strip():
                st.error("Please provide a response before proceeding.")
            else:
                # Save the response
                st.session_state.user_responses[current_question] = user_response
                
                if st.session_state.current_question_index < len(questions) - 1:
                    # Move to the next question
                    st.session_state.current_question_index += 1
                else:
                    # Process all responses when on the last question
                    user_responses_text = st.session_state.user_responses  # Pass the dictionary directly

                    aws_output = process_text_and_user_input(user_responses_text)
                    st.session_state.final_document = aws_output
                    
                    st.write("**Technical Document:**")
                    st.write(aws_output)
                    
                    # Create a Word document with only the AWS output
                    doc = Document()
                    doc.add_heading('AWS Cloud Expert Analysis Output', 0)
                    doc.add_paragraph(aws_output)
                    
                    # Save the document to a BytesIO buffer
                    st.session_state.doc_buffer = io.BytesIO()
                    doc.save(st.session_state.doc_buffer)
                    st.session_state.doc_buffer.seek(0)
                    
                    # Reset the session state for the next user
                    st.session_state.current_question_index = 0
                    st.session_state.user_responses = {}

# Provide download button if the document has been generated
if st.session_state.doc_buffer:
    st.download_button(
        label="Download AWS Output Document",
        data=st.session_state.doc_buffer,
        file_name='aws_output.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )









